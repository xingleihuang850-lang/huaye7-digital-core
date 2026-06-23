#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Word 数据说明书 + 掌握指南：deliverables/GJ5-15_数据说明书与掌握指南.docx
数据来源：experiments/master_stats.json（ETL 管线产出）。
用法： .venv/bin/python src/make_word.py
"""
import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ST = json.load(open(os.path.join(ROOT, "experiments", "master_stats.json"), encoding="utf-8"))
FIG = os.path.join(ROOT, "experiments", "figures")
OUT = os.path.join(ROOT, "deliverables", "GJ5-15_数据说明书与掌握指南.docx")
META = ST["curve_meta"]; SEGS = ST["segments"]; CURVE_ORDER = ST["curve_order"]

LATIN = "Arial"; CJK = "PingFang SC"
NAVY = RGBColor(0x1F, 0x38, 0x64); BLUE = RGBColor(0x2E, 0x54, 0x96)
GREY = RGBColor(0x80, 0x80, 0x80)


def set_cjk(run):
    run.font.name = LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = LATIN; st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    for i, sz in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = LATIN; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = NAVY if i == 1 else BLUE
        h.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def para(doc, text="", size=11, bold=False, color=None, align=None, after=6, before=0, italic=False):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
        set_cjk(r)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; set_cjk(r)
    r = p.add_run(text); set_cjk(r)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; set_cjk(r)
    r = p.add_run(text); set_cjk(r)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    t = OxmlElement("w:t"); t.text = "（在 Word 中右键 → 更新域，可生成目录）"
    fld.append(t); run._element.addprevious(fld)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)


def make_table(doc, headers, rows, widths=None, head_fill="2E5496", zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = ""
        r = hc[i].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); set_cjk(r)
        hc[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(hc[i], head_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(str(val)); r.font.size = Pt(9.5); set_cjk(r)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if zebra and ri % 2 == 1: shade(cells[ci], "EEF3FA")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def add_fig(doc, fname, width_in=6.3, caption=None):
    p = os.path.join(FIG, fname)
    if os.path.exists(p):
        doc.add_picture(p, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            para(doc, caption, size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.2); sec.left_margin = sec.right_margin = Cm(2.4)
    style_base(doc)

    # 页脚页码
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(); fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._element.append(fld); run.font.size = Pt(9)

    # ---------- 封面 ----------
    for _ in range(3): para(doc)
    para(doc, "GJ5-15 数字井筒", size=30, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, "数据说明书 与 掌握指南", size=22, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    para(doc, "Generative Digital Wellbore — Data Manual & Mastery Guide",
         size=12, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=40)
    para(doc, f"井段 {ST['interval_m'][0]}–{ST['interval_m'][1]} m   ·   连续曲线 {ST['n_samples']:,} 点/条   ·   5 个深度段",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    para(doc, "面向：数据掌握 / 答辩问询 / 实验设计", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, "生成日期 2026-06-22", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------- 目录 ----------
    para(doc, "目录", size=16, bold=True, color=NAVY, after=8)
    add_toc(doc)
    doc.add_page_break()

    # ===== 第1章 总览 =====
    doc.add_heading("1  这批数据是什么（一句话先记住）", level=1)
    para(doc, "一句话：这是吉林大学 GJ5-15 井一段 4.4 米岩心、经 CT 扫描 + PerGeos 相分割后反演出的"
              "“数字井筒”——把零散的几块岩心，变成了沿井深近乎连续的多条定量曲线（孔隙度、含油、泥质、裂缝等）"
              "加上分段统计、渗透率、储层评价。它已经是“判别式分割”的成品；本研究要做的下一跳是“生成式”补全与外推。",
         after=8)
    para(doc, "关键事实速记", size=12, bold=True, color=BLUE, after=4)
    facts = [
        ("井号 / 样品", ST["well"]),
        ("岩心井段", f"{ST['interval_m'][0]} – {ST['interval_m'][1]} m，共 {ST['interval_length_m']} m"),
        ("深度分段", "5 段：" + " / ".join(SEGS)),
        ("连续曲线", f"9 条，各 {ST['n_samples']:,} 个采样点，垂向分辨率约 {ST['vertical_resolution_mm']} mm"),
        ("CT 规格", f"像素 {ST['ct_spec']['pixel_um']} μm，{ST['ct_spec']['image']}，{ST['ct_spec']['bit_depth']}，{ST['ct_spec']['format']}"),
        ("2D 切片覆盖", "仅前 3 段有原始切片（3399 / 4235 / 4253 张）；后 2 段无 2D 切片"),
        ("缺口（=机会）", "无矿物种类标注、无体素级 Ground Truth、后 2 段缺原始切片"),
    ]
    make_table(doc, ["项目", "内容"], facts, widths=[4.5, 12.5])
    para(doc, after=4)
    para(doc, "数据可信度：本文所有数字由 src/verify_integrity.py 独立核验（15 项全过）——直接从原始逐片 "
              "txt/xlsx 重算，分段【均值】与服务商汇总表自报值逐项吻合（最大偏差 0.02）。文中 min/max 为逐片"
              "原始极值，服务商汇总表按深度聚合后报、极差略窄，二者均值一致，属口径差异而非错误。",
         size=9, italic=True, color=GREY, after=6)

    # ===== 第2章 数据怎么组织 =====
    doc.add_heading("2  数据怎么组织、存在哪", level=1)
    para(doc, "原始数据放在可移动外置盘（拔盘后软链接会失效，但已入 DuckDB 的结构化数据不受影响）。"
              "工程根用软链接引用，不拷贝。结构化数据统一进 DuckDB 仓库，agent 走 SQL 读取，不直接 glob 裸文件。", after=6)
    make_table(doc, ["层 / 位置", "内容", "读取方式"], [
        ["data/raw_stats/综合柱状图数据/*.txt", "9 条连续曲线 + 储层评价 + 渗透率（GBK 编码）", "ETL 解析 → DuckDB"],
        ["data/raw_stats/<属性>/*.xlsx", "各属性逐段逐层明细（含相体积、岩心体积列）", "ETL 解析 → DuckDB"],
        ["孔隙直径/…柱状饼状图.xlsx", "孔隙等效直径分级分布", "ETL 解析 → DuckDB"],
        ["GJ5-15_数据分析汇总_*.xlsx", "数据清单 + 分段统计 + 可行性分析", "人读 / 参考"],
        ["GJ5-15-图片库/", "CT 2D 切片、3D 渲染、剖面展开图", "LanceDB 多模态索引（计划）"],
        ["data/warehouse.db", "DuckDB 结构化仓库（本手册所有数字的来源）", "SQL"],
    ], widths=[6.0, 7.5, 3.5])
    para(doc, after=2)
    para(doc, "DuckDB 仓库内的表：curves（连续曲线宽表）、curves_long（长表）、segment_stats（分段统计）、"
              "permeability（渗透率）、pore_diameter（孔径）、reservoir_grade（储层评价）。"
              "重建命令：.venv/bin/python src/etl_build_warehouse.py", size=10, color=GREY, after=6)

    # ===== 第3章 每个参数怎么读 =====
    doc.add_heading("3  九条曲线，每条怎么读", level=1)
    para(doc, "下表给出全井范围与一句话物理含义。记住单位都是体积百分比（%），按逐层（沿井深每一薄片）统计。", after=6)
    desc = {
        "matrix_porosity":"基质（粒间/粒内）孔隙，储集空间主体，越高越好",
        "total_porosity":"基质+孔洞缝总孔隙，整体储集能力",
        "fracture_porosity":"裂缝贡献的孔隙，反映裂缝发育——本井是“渗流高速路”的指标",
        "total_vug_frac":"孔、洞、缝综合比率",
        "vug_fill_rate":"孔洞缝被充填的比率，越高有效连通越差",
        "oil_content":"含油孔隙体积 / 岩心体积",
        "oil_saturation":"含油孔隙体积 / 总孔隙体积，最关键的含油性指标",
        "shale_content":"泥质（黏土）含量 Vsh，越高物性通常越差",
        "diagenetic_min":"成岩自生矿物，本井极低（均值 0.04%），近似可忽略",
    }
    rows = []
    for k in CURVE_ORDER:
        o = ST["overall"][k]
        rows.append([META[k]["zh"], f"{o['mean']:.2f}", f"{o['min']:.2f}–{o['max']:.2f}", desc[k]])
    make_table(doc, ["曲线", "全井均值%", "范围%", "怎么读"], rows, widths=[2.6, 2.2, 2.6, 9.6])
    para(doc, after=6)
    add_fig(doc, "01_well_tracks.png", width_in=6.2,
            caption="图 1  数字井筒综合测井道——浅色为逐层原始值，深色为平滑趋势，虚线为段界")

    # ===== 第4章 数据告诉我们什么 =====
    doc.add_heading("4  数据讲了一个什么地质故事", level=1)
    seg4 = SEGS[3]; ps = ST["per_segment"]
    para(doc, "把 5 段横向一比，故事很清楚：", after=4)
    bullet(doc, f"第 4 段（{seg4}）是“甜点”。它同时拥有全井最高的含油率"
                f"（{ps[seg4]['oil_content']['mean']:.2f}%）、含油饱和度（{ps[seg4]['oil_saturation']['mean']:.2f}%）、"
                f"裂缝孔隙度（{ps[seg4]['fracture_porosity']['mean']:.2f}%）和最低的泥质含量"
                f"（{ps[seg4]['shale_content']['mean']:.2f}%），渗透率 Kh 也在此段冲到 87.8 mD 的峰值。",
           bold_prefix="储集甜点：")
    sh0 = ps[SEGS[0]]['shale_content']['mean']; sh4 = ps[SEGS[4]]['shale_content']['mean']
    bullet(doc, f"泥质含量自上而下递减（{sh0:.1f}% → {sh4:.1f}%），与含油性此消彼长——"
                "泥质越低、物性越好、含油越高，是一条贯穿全井的主线。", bold_prefix="泥质梯度：")
    bullet(doc, "孔隙结构上，小孔（50–100μm）数量占约 92%，但大孔（>100μm）贡献了约 40% 的孔隙体积——"
                "储集空间由少数大孔主导，这对建模时的采样与超分辨很关键。", bold_prefix="孔隙结构：")
    bullet(doc, "渗透率各向异性 Kh/Kv 在第 4 段最高（6.59），说明该段层理/裂缝定向性最强，渗流方向性明显。",
           bold_prefix="渗流方向：")
    para(doc, after=4)
    add_fig(doc, "02_segment_bars.png", width_in=6.3,
            caption="图 2  各深度段关键参数均值对比——第 4 段在含油、裂缝、充填上集中拔高")
    # 相关性要点
    para(doc, "相关性要点（逐层 Pearson）", size=12, bold=True, color=BLUE, before=6, after=4)
    corr = ST["correlation"]
    pairs = [("oil_content", "oil_saturation"), ("shale_content", "oil_saturation"),
             ("matrix_porosity", "total_porosity"), ("fracture_porosity", "total_vug_frac")]
    for a, b in pairs:
        r = corr[a][b]
        sign = "正相关" if r > 0 else "负相关"
        bullet(doc, f"{META[a]['zh']} 与 {META[b]['zh']}：r = {r:.2f}（{sign}）")

    # ===== 第5章 生成式 =====
    doc.add_heading("5  生成式数字井筒：缺口在哪、机会在哪", level=1)
    para(doc, "现状停在“判别式分割”：服务商已用 PerGeos 把相（基质/孔隙/裂缝）分好，并给出逐层定量统计。"
              "但要从“几块扫描过的岩心”推到“一整条数字井筒”，还差“生成”这一跳。本数据集天然暴露了三个缺口，"
              "正好是生成式建模的发力点：", after=6)
    make_table(doc, ["缺口", "现状", "生成式怎么补"], [
        ["后 2 段无 2D 切片", "3441.70–3443.52 m 仅有统计与 3D 渲染，无原始切片", "用前 3 段切片训练，沿井深生成/外推缺失段虚拟岩心"],
        ["无矿物种类标注", "只有相分割，未区分石英/长石/方解石/黏土", "2D（SEM-EDS）→3D（CT）迁移学习生成矿物分布"],
        ["分辨率有限 50μm", "难分辨 <100μm 颗粒", "三维体素超分辨，生成更细结构"],
        ["单点确定性结果", "每深度一个值，无不确定性", "多样本生成，给分布而非单点，做不确定性量化"],
    ], widths=[3.6, 6.4, 7.0])
    para(doc, after=4)
    para(doc, "技术路线（可行性表里已点名最匹配的一条）：2D→3D 迁移学习最贴合现有条件；"
              "配套可做体素超分辨（VAE/扩散）、坏井段补全（inpainting）、沿井深虚拟岩心生成（Transformer/扩散）。"
              "短板是缺体素级 Ground Truth 与矿物标签，需 SEM-EDS / XRD 补充。", after=6)

    # ===== 第6章 如何熟练掌握 =====
    doc.add_heading("6  如何熟练掌握这批数据（学习路径）", level=1)
    para(doc, "下面是一条“从看懂到能讲、能用于实验”的递进路径。建议照做一遍。", after=6)

    para(doc, "第一步：记住骨架数字（10 分钟）", size=12, bold=True, color=BLUE, after=3)
    bullet(doc, "井段 3439.13–3443.52 m，5 段，9 条曲线各约 1.8 万点。")
    bullet(doc, "甜点在第 4 段；泥质自上而下递减；成岩矿物几乎为 0。")
    bullet(doc, "全井含油饱和度均值约 30%，基质孔隙度约 16%，泥质约 26%。")

    para(doc, "第二步：能用 SQL 自己查（动手）", size=12, bold=True, color=BLUE, before=6, after=3)
    para(doc, "打开仓库：duckdb data/warehouse.db，几条最常用查询：", size=10, after=3)
    for q in [
        "-- 各段含油饱和度均值\nSELECT segment_label, AVG(oil_saturation) FROM curves GROUP BY 1 ORDER BY 1;",
        "-- 某深度区间的全部曲线\nSELECT * FROM curves WHERE depth_m BETWEEN 3441.7 AND 3442.0;",
        "-- 分段统计表直接读\nSELECT * FROM segment_stats WHERE property='oil_saturation';",
    ]:
        p = doc.add_paragraph(); r = p.add_run(q); r.font.name = "Menlo"; r.font.size = Pt(9)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
        p.paragraph_format.space_after = Pt(3)

    para(doc, "第三步：能讲出故事（对外答辩）", size=12, bold=True, color=BLUE, before=6, after=3)
    bullet(doc, "用“甜点段 + 泥质梯度 + 孔隙由大孔主导”三句话概括地质特征。")
    bullet(doc, "用“判别式已完成、生成式是下一跳”一句话定位研究贡献。")
    bullet(doc, "随手能指出三个缺口（后两段无切片 / 无矿物标签 / 无体素 GT）。")

    para(doc, "第四步：能用于实验（做研究）", size=12, bold=True, color=BLUE, before=6, after=3)
    bullet(doc, "把连续曲线当作沿井深的 1D 序列，做补全/外推的 baseline（先线性插值，再上生成模型）。")
    bullet(doc, "前 3 段切片 → 训练；后 2 段 → 作为“需生成”的目标，天然的留出验证。")
    bullet(doc, "用相关性矩阵选条件变量（如用泥质/孔隙度作为条件，生成含油饱和度）。")

    para(doc, "常见问答演练（别人一问就能答）", size=12, bold=True, color=BLUE, before=6, after=3)
    qa = [
        ("这数据多少米、多少点？", "井段 4.4 m（3439.13–3443.52），9 条曲线各约 17,982 点，垂向约 0.24 mm 一个点。"),
        ("最好的层段是哪段？", f"第 4 段 {SEGS[3]}：含油、含油饱和度、裂缝孔隙度全井最高，泥质最低，Kh 峰值 87.8 mD。"),
        ("为什么叫“生成式”，不是普通数字岩心？", "现有是判别式相分割的成品；贡献点是用生成模型补全缺失段、做超分辨与不确定性量化。"),
        ("数据有什么短板？", "后两段无原始 2D 切片、无矿物种类标注、无体素级 Ground Truth；需 SEM-EDS/XRD 补充。"),
        ("孔隙以大孔还是小孔为主？", "数量上小孔（50–100μm）占约 92%，但体积上大孔贡献约 40%，储集靠少数大孔。"),
    ]
    make_table(doc, ["别人会问", "你这样答"], qa, widths=[5.0, 12.0])

    # ===== 附录 =====
    doc.add_page_break()
    doc.add_heading("附录 A  关键数字速查表", level=1)
    rows = []
    for k in CURVE_ORDER:
        o = ST["overall"][k]
        rows.append([META[k]["zh"], f"{o['mean']:.2f}", f"{o['median']:.2f}", f"{o['min']:.2f}", f"{o['max']:.2f}", f"{o['std']:.2f}"])
    make_table(doc, ["曲线", "均值", "中位", "最小", "最大", "标准差"], rows, widths=[3.0, 2.4, 2.4, 2.4, 2.4, 2.4])

    para(doc, after=4)
    doc.add_heading("附录 B  分段渗透率与储层评价", level=1)
    prows = [[p["segment_label"], f"{p['kh_mD']:.1f}", f"{p['kv_mD']:.1f}", f"{p['kh_kv_ratio']:.2f}"] for p in ST["permeability"]]
    make_table(doc, ["深度段", "Kh (mD)", "Kv (mD)", "Kh/Kv"], prows, widths=[5.0, 3.0, 3.0, 3.0])
    para(doc, after=4)
    gt = ST["reservoir_grade_thickness"]
    grows = [["a 优", f"{gt.get('a',0):.2f}"], ["b 良", f"{gt.get('b',0):.2f}"], ["c 中", f"{gt.get('c',0):.2f}"], ["d 差", f"{gt.get('d',0):.2f}"]]
    para(doc, "各评价等级累计厚度 (m)：", size=11, bold=True, after=3)
    make_table(doc, ["等级", "厚度 (m)"], grows, widths=[4.0, 4.0])

    para(doc, after=6)
    doc.add_heading("附录 C  术语中英对照", level=1)
    terms = [["含油饱和度", "Oil saturation / So"], ["含油率", "Oil content / HYL"],
             ["孔隙度", "Porosity"], ["泥质含量", "Vsh / NZ"], ["成岩矿物", "Diagenetic mineral / CYKW"],
             ["基质", "Matrix / KX"], ["裂缝", "Fracture / LF"], ["孔洞缝", "Vug-fracture / KDF"],
             ["渗透率", "Permeability (Kh, Kv)"], ["储层评价", "Reservoir grade (a/b/c/d)"]]
    make_table(doc, ["中文", "英文 / 缩写"], terms, widths=[5.0, 7.0])

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(">> 已生成", os.path.relpath(OUT, ROOT))


if __name__ == "__main__":
    build()
