#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""花页7 Word 说明书+掌握指南 → deliverables/花页7/花页7_数据说明书与掌握指南.docx"""
import os, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ST = json.load(open(os.path.join(ROOT,"experiments","hy7_stats.json"),encoding="utf-8"))
FIG = os.path.join(ROOT,"experiments","hy7_figures")
OUT = os.path.join(ROOT,"deliverables","花页7","花页7_数据说明书与掌握指南.docx")
LATIN="Arial"; CJK="PingFang SC"
NAVY=RGBColor(0x10,0x23,0x3A); BLUE=RGBColor(0x1C,0x72,0x93); GREY=RGBColor(0x80,0x80,0x80)

def cjk(run): run.font.name=LATIN; run._element.rPr.rFonts.set(qn("w:eastAsia"),CJK)
def base(doc):
    s=doc.styles["Normal"]; s.font.name=LATIN; s.font.size=Pt(11)
    s.element.rPr.rFonts.set(qn("w:eastAsia"),CJK)
    for i,sz in [(1,18),(2,14),(3,12)]:
        h=doc.styles[f"Heading {i}"]; h.font.name=LATIN; h.font.size=Pt(sz); h.font.bold=True
        h.font.color.rgb=NAVY if i==1 else BLUE; h.element.rPr.rFonts.set(qn("w:eastAsia"),CJK)
def para(doc,t="",size=11,bold=False,color=None,align=None,after=6,italic=False,before=0):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if t:
        r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic
        if color: r.font.color.rgb=color
        cjk(r)
    return p
def bullet(doc,t,prefix=None):
    p=doc.add_paragraph(style="List Bullet")
    if prefix:
        r=p.add_run(prefix); r.bold=True; cjk(r)
    r=p.add_run(t); cjk(r); p.paragraph_format.space_after=Pt(3); return p
def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hx); tcPr.append(sh)
def table(doc,headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10)
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); cjk(r); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        shade(c,"1C7293")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,v in enumerate(row):
            cells[ci].text=""; r=cells[ci].paragraphs[0].add_run(str(v)); r.font.size=Pt(9.5); cjk(r)
            cells[ci].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT
            if ri%2==1: shade(cells[ci],"EEF3FA")
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    return t
def fig(doc,fn,w=6.3,cap=None):
    p=os.path.join(FIG,fn)
    if os.path.exists(p):
        doc.add_picture(p,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        if cap: para(doc,cap,size=9,italic=True,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER,after=10)
TOC_ENTRIES=[
    ("1  这是什么数据（一句话先记住）","3"),
    ("2  六个成像尺度，每个怎么读","4"),
    ("3  多尺度孔隙度：为什么不单调","5"),
    ("4  矿物组成：一块什么样的页岩","6"),
    ("5  孔喉结构与连通性","7"),
    ("6  核心任务与缺口（你研究的发力点）","8"),
    ("7  如何熟练掌握这套数据（学习路径）","8"),
    ("附录  多尺度关键数字速查","10"),
]
def toc(doc, content_width):
    """手动目录：标题左、页码右对齐、点线连，单行。避免自动域在 Pages 里页码换行。"""
    for title,page in TOC_ENTRIES:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10)
        p.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r=p.add_run(title); r.font.size=Pt(12.5); cjk(r)
        r2=p.add_run("\t"+page); r2.font.size=Pt(12.5); cjk(r2)

def modal_bin(scale_key):
    pr=ST["scale_data"][scale_key]["pore_radius"]
    return max(pr,key=lambda d:d["v1"]) if pr else None

def build():
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=sec.bottom_margin=Cm(2.2); sec.left_margin=sec.right_margin=Cm(2.4)
    base(doc)
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=f.add_run(); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); run._element.append(fld); run.font.size=Pt(9)

    # 封面
    for _ in range(3): para(doc)
    para(doc,"花页7井 多尺度数字岩心",size=30,bold=True,color=NAVY,align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
    para(doc,"数据说明书 与 掌握指南",size=22,bold=True,color=BLUE,align=WD_ALIGN_PARAGRAPH.CENTER,after=20)
    para(doc,"HuaYe-7 Multi-scale Digital Rock — Data Manual & Mastery Guide",size=12,italic=True,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER,after=40)
    para(doc,"深度 4199.21 m  ·  页岩  ·  6 个成像尺度（25μm → 8.589nm）",size=12,align=WD_ALIGN_PARAGRAPH.CENTER,after=6)
    para(doc,"面向：数据掌握 / 多尺度对齐 / 矿物分割 / 答辩问询",size=11,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
    para(doc,"数据已核校（src/verify_hy7.py，12 项全过）· 生成 2026-06-22",size=10,color=GREY,align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()
    para(doc,"目录",size=16,bold=True,color=NAVY,after=12)
    toc(doc, sec.page_width - sec.left_margin - sec.right_margin); doc.add_page_break()

    # 1
    doc.add_heading("1  这是什么数据（一句话先记住）",level=1)
    para(doc,"一句话：这是吉林大学花页7井 4199.21 m 一块页岩样品，被从 25μm 一路扫到 8.589nm 共 6 个成像尺度，"
             "服务商（欧勒姆能源测试）已做完各尺度的孔隙/喉道/裂缝三维分割与矿物定量。它是一套现成的「多尺度数字岩心」，"
             "研究要做的是把这些尺度对齐、融合，并用生成式方法补全与外推。",after=8)
    para(doc,"关键事实速记",size=12,bold=True,color=BLUE,after=4)
    g=ST["lithology_groups"]
    table(doc,["项目","内容"],[
        ["井 / 深度 / 岩性","花页7井 / 4199.21 m / 页岩（钙质长英质）"],
        ["成像尺度","6 个：25μm Amics 矿物 → 14μm/2.8μm 微米CT → 65nm 纳米CT → 200/10nm Maps SEM → 8.589nm FIB-SEM"],
        ["跨尺度倍率","分辨率从 25μm 到 8.589nm，跨约 3000 倍"],
        ["多尺度总孔隙度","1.56% / 7.18% / 1.93% / 0.41%（14μm/2.8μm/65nm/Maps，非单调，见第 3 章）"],
        ["矿物组成","长英质≈{}% + 碳酸盐≈{}% + 黏土≈{}%，黄铁矿{}%".format(g['长英质'],g['碳酸盐'],g['黏土'],g['黄铁矿'])],
        ["FIB-SEM","孔隙 {:,} 个 / 喉道 {:,} 个；有机质占比 {}%（注：非孔隙度）".format(ST['fib_counts']['pores'],ST['fib_counts']['throats'],ST['fib_organic_pct'])],
        ["数据可信度","各尺度统计由源文件独立核校（src/verify_hy7.py 12 项全过）"],
    ],widths=[4.0,13.0])

    # 2
    doc.add_heading("2  六个成像尺度，每个怎么读",level=1)
    rows=[]
    for s in ST["scales"]:
        rows.append([s["name"],s["res"],s["modality"],s["content"]])
    table(doc,["尺度","分辨率","模态","主要内容"],rows,widths=[3.0,2.6,3.4,8.0])
    para(doc,after=4)
    fig(doc,"06_ladder.png",6.4,"图 1  多尺度成像阶梯：从矿物(25μm)到纳米孔(8.589nm)，跨约 3000 倍")
    para(doc,"读法：尺度越细，看见的孔越小、数量越多但单孔体积越小。不同尺度看的是同一块岩石的不同层级孔隙系统，"
             "因此孔隙度不会简单相加，需要「多尺度对齐 + 融合」才能拼出完整图景。",after=6)

    # 3
    doc.add_heading("3  多尺度总孔隙度：为什么不单调",level=1)
    fig(doc,"01_porosity.png",6.2,"图 2  各尺度总孔隙度（4 个尺度）——2.8μm 最高，呈非单调")
    para(doc,"原始扫描报告均给出“总孔隙度”：14μm=1.558%、2.8μm=7.182%、65nm=1.925%、Maps(10nm)=0.411%。"
             "（其中 14μm/2.8μm/65nm 已由各尺度逐切片面孔率均值独立复现。）"
             "为什么 2.8μm 反而最高？因为不同尺度的视场、可分辨孔径、有效统计体不同——这正是「尺度效应 / 代表性体元(REV)」，"
             "是页岩多尺度表征的核心讨论点，不是数据矛盾。",after=4)
    bullet(doc,"14μm 与 65nm 偏低：分辨率限制 + 视场内有效孔隙比例。",prefix="解读：")
    bullet(doc,"2.8μm 偏高：恰好捕捉到微孔-微裂缝主峰，且统计体足够。")
    bullet(doc,"汇报时把它当作「需要多尺度融合」的证据，而不是误差。")
    para(doc,"⚠️ 注意：FIB-SEM 报告给的 1.52% 是“有机质占比”，不是孔隙度，故未计入上面的总孔隙度对比；"
             "FIB-SEM 本身只直接报了孔隙/喉道/裂缝的体积与计数，未单列总孔隙度。",size=10,color=GREY,after=6)

    # 4
    doc.add_heading("4  矿物组成：一块什么样的页岩",level=1)
    fig(doc,"02_minerals.png",6.4,"图 3  矿物组成（Amics 粗扫 25μm）与岩性归组")
    para(doc,"按 Amics 粗扫(25μm)：长英质（斜长石 30.9% + 石英 23.5% + 钾长石）约 {}%，"
             "碳酸盐（方解石 18.1% + 白云石 12.3% + 铁白云石 10.9%）约 {}%，"
             "黏土（绿泥石+伊利石等）仅约 {}%，黄铁矿 {}%。属钙质长英质页岩、黏土含量低。"
             .format(g['长英质'],g['碳酸盐'],g['黏土'],g['黄铁矿']),after=4)
    para(doc,"⚠️ 粗扫(25μm)与精扫(1μm)矿物比例不同（精扫方解石↑到 29%、石英↓到 17%、黏土↑）。"
             "做矿物分割时要先定清楚以哪一套为标定基准——这是视场代表性差异。",size=10,color=GREY,after=6)

    # 5
    doc.add_heading("5  孔喉结构与连通性",level=1)
    fig(doc,"03_pore_radius.png",6.3,"图 4  各尺度孔隙半径分布（体积 vs 个数）")
    rows=[]
    for k,nm in [("ct14","微米CT 14μm"),("ct28","微米CT 2.8μm"),("nano65","纳米CT 65nm"),("fib","FIB-SEM")]:
        mb=modal_bin(k)
        rows.append([nm, mb["bin"] if mb else "-", f"{mb['v1']}%" if mb else "-"])
    table(doc,["尺度","孔隙半径主峰区间","该区间体积占比"],rows,widths=[4.0,5.5,4.0])
    para(doc,after=4)
    fig(doc,"04_coordination.png",6.0,"图 5  各尺度孔隙配位数分布（连通性）")
    para(doc,"各尺度配位数都集中在低值、0 为最大占比（14μm 高达 85%，2.8μm/65nm/FIB-SEM 约 36–47%），"
             "说明孔隙连通性整体偏弱、孤立孔较多——这对渗流与建模都很关键。",after=6)

    # 6
    doc.add_heading("6  核心任务与缺口（你研究的发力点）",level=1)
    table(doc,["任务","现状 / 缺口","怎么做"],[
        ["多尺度对齐(配准)","各尺度图像-标签未配准是真瓶颈","Plan B：直接用天然对齐的 sus.raw(灰度体)+pore.raw(标签)训练，绕开原图配准"],
        ["矿物分割","已有 Amics 矿物图，但粗/精扫基准不一","以某一尺度定标定基准，迁移到 CT 体；长英质/碳酸盐/黏土分类"],
        ["多尺度融合","5 个尺度孔隙度非单调、各看一层","跨尺度孔隙网络拼接 + REV 分析"],
        ["生成式补全","单点统计、缺连续体","体素超分辨 / 跨尺度生成 / 不确定性量化"],
    ],widths=[3.0,6.0,8.0])

    # 7
    doc.add_heading("7  如何熟练掌握这套数据（学习路径）",level=1)
    para(doc,"第一步：记住骨架（10 分钟）",size=12,bold=True,color=BLUE,after=3)
    bullet(doc,"花页7 @ 4199.21m 页岩；6 尺度 25μm→8.589nm；跨约 3000 倍。")
    bullet(doc,"多尺度孔隙度非单调（2.8μm 最高 7.18%）= 尺度效应/REV。")
    bullet(doc,"钙质长英质页岩：长英质55%+碳酸盐41%+黏土3%；FIB-SEM 7.6万孔/6.2万喉道。")
    para(doc,"第二步：能讲出每个尺度看什么（对外答辩）",size=12,bold=True,color=BLUE,before=6,after=3)
    bullet(doc,"Amics=矿物；微米CT=骨架孔/裂缝；纳米CT/Maps/FIB-SEM=纳米孔与有机孔。")
    bullet(doc,"一句话定位贡献：数据现成、阶段一同体内部基线已建立；空间独立验证待未来重跑，下一跳是多尺度对齐+生成式融合。")
    para(doc,"第三步：能用于实验（做研究）",size=12,bold=True,color=BLUE,before=6,after=3)
    bullet(doc,"先跑 Plan B（sus.raw+pore.raw 天然对齐）做三维分割 baseline。")
    bullet(doc,"用矿物图(Amics)做矿物分割标签，注意粗/精扫基准。")
    bullet(doc,"用多尺度孔隙度差异做 REV / 尺度融合的讨论与建模条件。")
    para(doc,"常见问答演练",size=12,bold=True,color=BLUE,before=6,after=3)
    table(doc,["别人会问","你这样答"],[
        ["这数据多少个尺度？","6 个，25μm Amics 到 8.589nm FIB-SEM，跨约 3000 倍，同一块 4199.21m 页岩。"],
        ["孔隙度为什么对不上？","尺度效应/REV——不同尺度看不同层级孔隙，2.8μm 最高 7.18%，需多尺度融合而非相加。"],
        ["这是什么岩石？","钙质长英质页岩：长英质≈55%、碳酸盐≈41%、黏土仅≈3%，黄铁矿 0.5%。"],
        ["最大难点是什么？","图像-标签配准；解法是用天然对齐的 sus.raw+pore.raw 训练(Plan B)。"],
        ["数据可信吗？","已独立核校 12 项全过，汇总与服务商原始统计逐项一致。"],
    ],widths=[5.0,12.0])

    # 附录
    doc.add_page_break()
    doc.add_heading("附录  多尺度关键数字速查",level=1)
    rows=[]
    for p in ST["porosity_summary"]:
        rows.append([p["scale"],str(p["res"]),p["metric"],f"{p['porosity']}%",str(p.get("note") or "")])
    table(doc,["尺度","分辨率","指标","数值","备注"],rows,widths=[3.0,2.2,2.8,2.0,7.0])
    para(doc,"数据来源：外置盘「吉林大学数据报告归总」各尺度子目录（服务商原始统计，只读）→ src/hy7_etl.py → hy7_stats.json。"
             "全部数字经 src/verify_hy7.py 核校（12 项全过）。",size=9,italic=True,color=GREY,after=4)

    os.makedirs(os.path.dirname(OUT),exist_ok=True); doc.save(OUT)
    print(">> 已生成",os.path.relpath(OUT,ROOT))

if __name__=="__main__": build()
